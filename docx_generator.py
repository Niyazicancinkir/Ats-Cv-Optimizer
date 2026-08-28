import os


def _add_list_paragraph(document, title, items):
    if not items:
        return
    document.add_heading(title, level=2)
    if isinstance(items, list):
        for item in items:
            document.add_paragraph(str(item), style="List Bullet")
    else:
        document.add_paragraph(str(items))


def generate_docx(cv_data, output_docx_path):
    try:
        from docx import Document
    except Exception:
        print("[!] DOCX olusturma icin 'python-docx' gerekli. Kurulum: pip install python-docx")
        return False

    try:
        os.makedirs(os.path.dirname(output_docx_path), exist_ok=True)
        document = Document()

        document.add_heading(cv_data.get("name", ""), level=1)
        title = cv_data.get("title", "")
        if title:
            document.add_paragraph(title)

        contact_parts = [
            cv_data.get("location", ""),
            cv_data.get("email", ""),
            cv_data.get("phone", ""),
            cv_data.get("linkedin", ""),
            cv_data.get("github", ""),
        ]
        contact_text = " | ".join(part for part in contact_parts if part)
        if contact_text:
            document.add_paragraph(contact_text)

        summary = cv_data.get("summary", "")
        if summary:
            document.add_heading("Summary", level=2)
            document.add_paragraph(summary)

        experience = cv_data.get("experience", [])
        if experience:
            document.add_heading("Work Experience", level=2)
            for job in experience:
                role = job.get("role", "")
                company = job.get("company", "")
                period = job.get("period", "")
                location = job.get("location", "")
                header = " - ".join(part for part in [role, company] if part)
                if header:
                    document.add_paragraph(header)
                meta = " | ".join(part for part in [period, location] if part)
                if meta:
                    document.add_paragraph(meta)
                description = job.get("description", "")
                if description:
                    document.add_paragraph(description, style="List Bullet")

        projects = cv_data.get("projects", [])
        if projects:
            document.add_heading("Projects", level=2)
            for project in projects:
                project_name = project.get("name", "")
                project_url = project.get("url", "")
                project_header = " - ".join(part for part in [project_name, project_url] if part)
                if project_header:
                    document.add_paragraph(project_header)
                project_description = project.get("description", "")
                if project_description:
                    document.add_paragraph(project_description, style="List Bullet")

        education = cv_data.get("education", [])
        if education:
            document.add_heading("Education", level=2)
            for edu in education:
                school = edu.get("school", "")
                degree = edu.get("degree", "")
                period = edu.get("period", "")
                gpa = edu.get("gpa", "")
                line = " | ".join(part for part in [school, degree, period, f"GPA: {gpa}" if gpa else ""] if part)
                if line:
                    document.add_paragraph(line)

        _add_list_paragraph(document, "Technical Skills", cv_data.get("skills", []))
        _add_list_paragraph(document, "Certifications", cv_data.get("certifications", []))

        document.save(output_docx_path)
        return True
    except Exception as e:
        print(f"[!] DOCX olusturulurken hata: {e}")
        return False